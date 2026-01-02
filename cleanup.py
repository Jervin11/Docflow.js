#cleanup.py
import pandas as pd
import io
import re

def read_docx_tables(file_bytes: bytes):
    """Extracts tables from a Word document (.docx)."""
    try:
        from docx import Document
    except Exception as e:
        raise ImportError("python-docx is required to read .docx files. Install with `pip install python-docx`") from e

    doc = Document(io.BytesIO(file_bytes))
    tables = []
    # --- Extract actual tables ---
    for table in doc.tables:
        data = []
        for row in table.rows:
            data.append([cell.text.strip() for cell in row.cells])
        if data:
            tables.append(pd.DataFrame(data))

    
    paragraph_data = extract_paragraph_data(doc)
    if paragraph_data is not None:
        tables.append(paragraph_data)

    return tables

def read_pdf_tables(file_bytes: bytes):
    """Extracts tables from a PDF file using pdfplumber.

    Also attempts to extract colon-separated paragraph records (Key: Value) from
    the page text and returns them as a DataFrame when present.
    """
    try:
        import pdfplumber
    except Exception as e:
        raise ImportError("pdfplumber is required to read .pdf files. Install with `pip install pdfplumber`") from e

    tables = []
    text_blocks = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            extracted = page.extract_table()
            if extracted:
                tables.append(pd.DataFrame(extracted))
            # collect text for paragraph parsing
            page_text = page.extract_text()
            if page_text:
                text_blocks.append(page_text)

    # try to parse paragraph-style key:value records from the text
    if text_blocks:
        pdf_para_df = extract_paragraph_data_from_text("\n".join(text_blocks))
        if pdf_para_df is not None:
            tables.append(pdf_para_df)

    return tables

def read_excel(file_bytes: bytes):
    """Reads an Excel file into DataFrames."""
    excel_data = pd.read_excel(io.BytesIO(file_bytes), sheet_name=None)
    return list(excel_data.values())







def extract_paragraph_data(doc):
    """Detect paragraph data like Name:Abi, Age:20, City:Madurai."""
    # Split paragraphs and keep even empty lines to detect record breaks
    paragraphs = [p.text.strip() for p in doc.paragraphs]
    lines = [p for p in paragraphs if p or p == ""]

    records = []
    current = {}

    # Detect available keys dynamically (preserve order of first appearance)
    key_pattern = re.compile(r"^([\w\s]+?)\s*:")
    all_keys_ordered = []

    # First pass – collect keys in first-seen order
    for line in lines:
        m = key_pattern.match(line)
        if m:
            k = m.group(1).strip()
            if k not in all_keys_ordered:
                all_keys_ordered.append(k)

    # Second pass – build records
    for line in lines:
        if not line.strip():  # blank line = end of one record
            if current:
                records.append(current)
                current = {}
            continue

        if ":" in line:
            k, v = line.split(":", 1)
            k, v = k.strip(), v.strip()
            current[k] = v

    if current:
        records.append(current)

    # If no structured key-value data found
    if not records or len(all_keys_ordered) < 1:
        return None

    # Keep only complete records where all keys are present and non-empty
    complete_records = []
    for r in records:
        if all((k in r and str(r[k]).strip() != "") for k in all_keys_ordered):
            complete_records.append(r)

    if not complete_records:
        return None

    # Convert to DataFrame using ordered columns
    df = pd.DataFrame(complete_records, columns=all_keys_ordered)
    # Title-case column names
    df.columns = [c.strip().title() for c in df.columns]
    return df
    ef clean_table(df: pd.DataFrame) -> pd.DataFrame:
    """Cleans up misaligned, merged, or blank rows without reordering columns."""
    df = df.dropna(how="all").reset_index(drop=True)

    # Shift non-empty cells left if there are blanks (count non-empty values correctly)
    def _shift_row(row):
        values = [x for x in row if (pd.notna(x) and str(x).strip() != "")]
        return pd.Series(values + [None] * (len(row) - len(values)))

    df = df.apply(_shift_row, axis=1)

    # Dynamically detect whether the first row is a header row.
    # Heuristics used:
    # - If a majority of first-row cells look like short alphabetic labels (no digits)
    #   and there are few or no numeric-looking tokens, treat as header.
    # - Also check column-wise: if first-row value in a column looks non-numeric but
    #   the next few rows in that column are mostly numeric, that's a strong signal
    #   the first row is a header for that column.
    def _looks_like_header_row(row, df):
        row_vals = [str(x).strip() for x in row.tolist()]
        non_empty = [v for v in row_vals if v and v.lower() not in ("nan", "none")]
        if len(non_empty) == 0:
            return False

        # tokens that look like header labels (mostly letters, spaces and punctuation)
        header_like = sum(1 for v in non_empty if re.match(r'^[A-Za-z\s\-\_/&()\.]+$', v) and len(v) <= 60)
        # tokens that look numeric
        numeric_like = sum(1 for v in non_empty if re.match(r'^-?\d+(?:\.\d+)?$', v))

        # If majority of non-empty cells look like textual labels and few numeric tokens
        if header_like >= max(1, len(df.columns) // 2) and numeric_like == 0:
            return True

        # Column-wise check: if first row is non-numeric and the rows beneath are mostly numeric
        col_votes = 0
        rows_to_check = min(5, max(0, len(df) - 1))
        if rows_to_check > 0:
            for i, v in enumerate(row_vals):
                if v == "":
                    continue
                is_header_token = not re.match(r'^-?\d+(?:\.\d+)?$', v)
                below = df.iloc[1:1 + rows_to_check, i].astype(str).str.strip().tolist()
                below_numeric = sum(1 for b in below if re.match(r'^-?\d+(?:\.\d+)?$', b))
                if is_header_token and below_numeric >= rows_to_check / 2:
                    col_votes += 1
            # if a reasonable fraction of columns suggest header->data relationship
            if col_votes >= max(1, len(df.columns) // 3):
                return True

        return False

    if _looks_like_header_row(df.iloc[0], df):
        # Use first row as header
        df.columns = [str(c).strip().title() for c in df.iloc[0]]
        data = df[1:].reset_index(drop=True)
    else:
        # No header detected — assign generic column names
        num_cols = df.shape[1]
        default_headers = [f"Col{i+1}" for i in range(num_cols)]
        df.columns = default_headers
        data = df.reset_index(drop=True)

    # Normalize empty strings to NA and strip strings
    def _clean_cell(x):
        if pd.isna(x):
            return pd.NA
        if isinstance(x, str):
            s = x.strip()
            return s if s != "" else pd.NA
        return x

    data = data.applymap(_clean_cell)

    # Only drop rows that are completely empty; keep rows with some missing cells
    data = data.dropna(how='all').reset_index(drop=True)

    return data
